import docx
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as plticker
import PySimpleGUI as sg
import os.path
import math
import datetime


def RemoveSomeFiles():
    os.remove('Rpm_Load.png')
    os.remove('load_thrrotle.png')
    os.remove('SOCvsDOD.png')
    os.remove('PowerDis.png')


def CreateGUI():
    x=0
    #window = sg.window("CSV Report Maker", layout)

def BaseCalc(DataLen):
    Base = (199900*(1/DataLen))/math.sqrt(DataLen*0.0008)
    return Base


def CreateDoc(name):
    doc = docx.Document()
    doc.add_heading('GENSET RUN Report - '+TestDate, 0)

    p = doc.add_paragraph('This file is a report of the data recorded during a certian run. The data referd to the the loaded file. The TILS Software is: '+ str(Tils_Software))
    doc.add_paragraph('Test Start Time: ' + str(Start_time), style='List Bullet')
    doc.add_paragraph('Test End Time: ' + str(End_time), style='List Bullet')
    try:
        Start_time_obj = datetime.datetime.strptime(Start_time, '%H:%M')                     ## Calculation way for haveing the test duration time
        End_time_obj = datetime.datetime.strptime(End_time, '%H:%M')
        def_time = End_time_obj - Start_time_obj
        doc.add_paragraph('Test Duration: ' + str(def_time), style='List Bullet')
    except:
        print('Problem with time column')

    #p = doc.add_paragraph('A plain paragraph having some ')
    #p.add_run('bold').bold = True
    #p.add_run(' and some ')
    #p.add_run('italic.').italic = True

    #doc.add_heading('Main data characters: RPM, LOAD, etc.', level=1)

    #############################    Main Characters     ###################################################
    doc.add_paragraph('Main Characters', style='Intense Quote')
    doc.add_paragraph('Max RPM value: '+str(max_rpm), style='List Bullet')
    rpm_test = ('Is grater then 2000? '+str(max_rpm>2000))
    doc.add_paragraph('Max Engine Temperature value: ' + str(max_EngTemp), style='List Bullet')
    doc.add_paragraph('Max Genrator Power value: ' + str(max_GenPower), style='List Bullet')
    doc.add_picture('Rpm_Load.png', width=Inches(6), height=Inches(4))

    doc.add_page_break()

    #########################    Other Data    ###############################################################
    doc.add_paragraph('Other Data', style='Intense Quote')
    #doc.add_paragraph('Max Bus DC voltage value: '+str(max_rpm), style='List Bullet')
    doc.add_paragraph('Engine Working time: '+str(WorkHours)+' Minutes', style='List Bullet')
    #doc.add_paragraph('Max RPM value: '+str(max_rpm), style='List Bullet')
    #doc.add_paragraph('Max RPM value: ' + str(max_rpm), style='List Bullet')
    doc.add_picture('load_thrrotle.png', width=Inches(6), height=Inches(4))
    doc.add_page_break()

    ##########################    Power and Electricty Data      ##################################################################
    doc.add_paragraph('Power and Electricty Data', style='Intense Quote')
    doc.add_paragraph('Max Bus DC voltage value: ' + str(max_BusV), style='List Bullet')
    doc.add_paragraph('Max Rectifire Temprature: ' + str(recti_T), style='List Bullet')
    doc.add_paragraph('Max PMG Temprature: ' + str(pmg_T), style='List Bullet')
    doc.add_picture('SOCvsDOD.png', width=Inches(6), height=Inches(4))
    doc.add_picture('PowerDis.png', width=Inches(6), height=Inches(4))
    #doc.add_picture('SOCvsDOD_1.png', width=Inches(6), height=Inches(4))

    #####################    Results    ######################################################################
    doc.add_paragraph('Results', style='Intense Quote')
    records = (
        ('Max RPM', str(max_rpm) , str(rpm_test)),
        #(7, '422', 'Eggs'),
        #(4, '631', 'Spam, spam, eggs, and spam')
    )


    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'
    hdr_cells[2].text = 'Pass/Fail'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    doc.add_page_break()

    File_name = ('GensetRunningReport_' +str(TestDate) + '.docx')
    print(File_name)
    #doc.save(os.path.join(File_name))
    #doc.save('GensetReport' + TestDate +'.docx')
    #doc.save('GensetRunningReport_' +str(TestDate) + '.docx')
    doc.save('GensetRunningReport_NoDate.docx')


def DataAnalyse(FileAdress):
    global DataLen, max_rpm, max_loadPower, max_GenPower, max_EngTemp, max_BusV, WorkHours, Tils_Software, TestDate, Start_time, End_time, recti_T, pmg_T
    Tils_Software = 0
    data = pd.read_csv(FileAdress, index_col=False) #with GUI it will be FileAdress
    DataLen = len(data)

    if DataLen > 4000:
        data = data.iloc[::3, :]
        DataLen = len(data)
    #print(data)

    DateList = []
    TestDate = data['Time'][3].split(' ', 1)[0]
    Start_time = data['Time'][3].split(' ', 1)[1]
    End_time = data['Time'][DataLen-2].split(' ', 1)[1]
    for x in data['Time']:
        DateList.append(x.split(' ', 1)[1])

    max_rpm = (data[' Rpm'].max())
    max_loadPower = (data[' LoadPower'].max())
    max_GenPower = (data[' GeneratorPower'].max())
    max_EngTemp = (data[' EngineTemperature'].max())
    max_dod = (data[' Dod'].max())
    max_BusV = (data[' BusDcVoltage'].max())
    WorkHours = data[' EngineWorkHours'].max()
    #Tils_Software = data[' SoftwareVersion'].max()
    recti_T = data[' RectifierTemperature'].max()
    pmg_T = data[' AlternatorTemperature'].max()

    fig, ax = plt.subplots()
    plt.xticks(rotation=90)
    plt.grid()
    plt.plot(DateList, data[' Rpm'], color='r', label='RPM')

    ax2 = ax.twinx()
    ax2.set_ylabel("GenPower", color="blue", fontsize=9)
    ax.set_ylabel("RPM", color='r', fontsize=9)
    ax2.plot(DateList, data[' GeneratorPower'], color='b', label='GenPower')
    loc = plticker.MultipleLocator(BaseCalc(DataLen))  # this locator puts ticks at the selected interval
    ax.xaxis.set_major_locator(loc)
    plt.savefig('Rpm_Load.png', dpi = 600, bbox_inches = 'tight')


    fig3, ax3 = plt.subplots()
    plt.xticks(rotation=90)
    plt.scatter(data[' Rpm'], data[' ECU_AirSystemThrottleValvePosition'], s=16, c=data[' LoadPower'])
    loc3 = plticker.MultipleLocator(100)
    ax3.xaxis.set_major_locator(loc3)
    ax3.grid(True)
    plt.title("RPM Vs TH.pos - Power distreubtion")
    plt.xlabel("RPM")
    plt.ylabel("Throttle position (%)")
    plt.colorbar()

    plt.legend(labels=data[' LoadPower'], title='Power')
    ax3.legend()
    plt.savefig('load_thrrotle.png', dpi = 600, bbox_inches = 'tight')

    fig5, ax5 = plt.subplots()
    plt.xticks(rotation=90)
    plt.plot(data['Time'], data[' BatteryPowerDuringCharge'], color='r')
    ax5.set_ylabel("Charge", color='r', fontsize=9)
    ax6 = ax5.twinx()
    ax6.plot(DateList, data[' BatteryPowerDuringDischarge'], color='b')
    ax6.set_ylabel("Discharge", color='b', fontsize=9)
    plt.legend()
    plt.title("Battery Power - Discharge & Charge")
    loc5 = plticker.MultipleLocator(BaseCalc(DataLen))
    ax5.xaxis.set_major_locator(loc5)
    plt.legend()
    plt.savefig('SOCvsDOD.png', dpi = 600, bbox_inches = 'tight')
    
    figB = plt.figure()
    axB = figB.add_axes([0, 0, 1, 1])
    axB.bar(data[' GeneratorPower'], data[' LoadPower'], 0.35, color='r')
    axB.bar(data[' GeneratorPower'], data[' BatteryPowerDuringCharge'], 0.35, bottom=data[' LoadPower'], color='b')
    axB.set_ylabel('Load&Batt Power')
    axB.set_xlabel('Gen Power')
    ax.set_xticks(np.arange(0, 3, 0.1))
    ax.set_yticks(np.arange(0, 3, 0.1))
    axB.legend(labels=['Load Power','Battery Power'])
    plt.savefig('PowerDis.png', dpi=600, bbox_inches='tight')
"""
    fig8, ax9 = plt.subplots()
    plt.xticks(rotation=30)
    plt.plot(data['Time'], data[' Dod'], color='r')
    ax9.set_ylabel("Charge", color='r', fontsize=9)
    ax10 = ax9.twinx()
    ax10.plot(data['Time'], data[' Soc'], color='b')
    ax10.set_ylabel("Discharge", color='b', fontsize=9)
    plt.legend()
    plt.title("Battery Power - Discharge & Charge")
    loc7 = plticker.MultipleLocator(base=10)
    ax9.xaxis.set_major_locator(loc7)
    plt.legend()
    plt.savefig('SOCvsDOD_1.png', dpi = 600, bbox_inches = 'tight')
    fig6, ax7 = plt.subplots()
    plt.xticks(rotation=30)
    #plt.plot(data['Time'], data[' '])
"""


def main():
    sg.theme("DarkTeal2")
    layout = [[sg.T(""), sg.Text('This software ganerete a report file from the EMS recordings.')],
              [sg.Text('The report will be created in the same folder of the selected data file.')],
              [sg.Text("Choose a Data File: "), sg.Input(key="-IN2-" ,change_submits=True), sg.FileBrowse(key="-IN-",)],
              [sg.Button("Submit")]]
    ###Building Window
    window = sg.Window('GENSET Report Ganerator', layout, size=(600, 150))
    while True:
        event, values = window.read()
        #print(values["-IN2-"])
        if event == sg.WIN_CLOSED or event == "Exit":
            break
        elif event == "Submit":
            try:
                DataAnalyse(values["-IN-"])
                CreateDoc('temp')
                RemoveSomeFiles()
                sg.Popup('Report file has created...', title='Succses!')
            except Exception as e:
                sg.Popup(e, title='Fail!')
            print('OK')
        print(values["-IN-"])

            #DataAnalyse('FileN')
    #CreateDoc('temp')


if __name__ == '__main__':
    main()